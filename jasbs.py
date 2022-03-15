import pytesseract
from PIL import Image

import cv2
import os
import shutil

from docx import Document
from docx.shared import Cm,Mm,Pt,Inches
# from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_LINE_SPACING


dir_path="D://covid-files//"
pytesseract.pytesseract.tesseract_cmd="C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

cnt=0
for img_file in os.listdir(dir_path):
    if not img_file.endswith(".jpeg"):
        continue
    img_path=dir_path+img_file
    # print(img_path)
    img=cv2.imread(img_path)
    gray=cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    gray = cv2.threshold (gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]

    gray = cv2.medianBlur (gray, 3)
    # choice=input("thresh or blur")
    # if choice=="thresh":
    #     gray = cv2.threshold (gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
    # elif choice=="blur":
    #     gray = cv2.medianBlur (gray, 3)

    gray_file_path=dir_path+"{}_tmp.jpg".format(img_file)
    cv2.imwrite(gray_file_path,gray)

    text=pytesseract.image_to_string(Image.open(gray_file_path))
    text=text.replace("\n"," ")
    # print(text)

    doc=Document()
    secs=doc.sections
    for sec in secs:
        # margins=0.3
        sec.top_margin=Inches(0.3)
        sec.bottom_margin=Inches(0.3)
        sec.left_margin=Inches(0.3)
        sec.right_margin=Inches(0.3)
        # A4
        sec.page_height=Mm(297)
        sec.page_width=Mm(210)

        paragraph=doc.add_paragraph()
        prg_fmt=paragraph.paragraph_format
        # 0pt vs 0 line.
        prg_fmt.space_before=Pt(0)
        prg_fmt.space_after=Pt(1.2)
        prg_fmt.line_spacing_rule=WD_LINE_SPACING.SINGLE

        run=paragraph.add_run(text)
        font=run.font
        font.name="Times New Roman"
        font.size=Pt(10)

    doc_file=img_file.replace(".jpeg",".doc")
    doc_path=dir_path+doc_file
    doc.save(doc_path)

    # print("one done.")
    cnt+=1
    os.remove(gray_file_path)

    new_dir=img_path.replace(".jpeg","")
    # print(new_dir)
    os.mkdir(new_dir)
    shutil.copy(img_path,new_dir+os.sep+img_file)
    shutil.copy(doc_path,new_dir+os.sep+doc_file)
    # print(new_dir+img_file)
    # print(new_dir+doc_file)
    print(cnt)
    # print("one done.")

    # if cnt>=2:
    #     break

        # sec.orientation=

    # cv2.imshow("img",img)
    # cv2.imshow("gray",gray)
    # cv2.waitKey(0)

